from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
MID_BLUE  = RGBColor(0x2B, 0x54, 0x7E)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def styled_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        shade_cell(c, '1B3A5C')
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        fill = 'EAF0F6' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            shade_cell(c, fill)
            c.text = ''
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(8.5)
    return t

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = DARK_BLUE
    return p

def p(text, bold=False, italic=False, size=10.5):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return para

def bullet(items):
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(item).font.size = Pt(10)

def code_block(text):
    para = doc.add_paragraph()
    para.style = doc.styles['No Spacing']
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)

# ─── COVER ───────────────────────────────────────────────────────────────────
doc.add_paragraph()
t = doc.add_paragraph('Blockchain Based Employment Platform\nfor Industrial and Service Workforce')
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in t.runs:
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = DARK_BLUE

doc.add_paragraph()
auth = doc.add_paragraph('Sargam Puram [B15]  |  Tejas Shinde [B28]  |  Kartik Shingde [B29]  |  Kranti Zagade [B44]\nGuide: Dr. Nikita Kulkarni')
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ─── ABSTRACT ────────────────────────────────────────────────────────────────
h('Abstract')
p('This project proposes a Blockchain Based Employment Platform for Industrial and Service Workforce to overcome challenges such as lack of transparency, delayed payments, and dependency on intermediaries. The system enables secure and direct interaction between workers and employers using blockchain technology. Smart contracts automate job agreements and payment processes, ensuring fairness and trust. The platform provides transparent, tamper-proof employment records on the Solana blockchain and improves efficiency in workforce management. The backend Node.js/Express server handles off-chain data while the React.js frontend provides an intuitive, Web3-enabled interface for workers, employers, and administrators.')

# ─── INTRODUCTION ────────────────────────────────────────────────────────────
h('1. Introduction')
p('Industrial and service workforce forms the backbone of economic growth, yet workers often face issues such as job insecurity, unfair wages, and delayed payments. Traditional employment systems are centralized and rely on intermediaries, which leads to lack of transparency and trust. Blockchain technology offers a decentralized and secure solution to these problems.')
p('This project, DeWages Network, designs a blockchain-based employment platform on the Solana network that connects workers and employers directly. A Rust smart contract (Anchor framework) acts as the trustless escrow and agreement engine while a Node.js backend and React frontend provide an accessible web-based experience.')

# ─── SCOPE AND METHODOLOGY ───────────────────────────────────────────────────
h('2. Scope and Methodology')
h('2.1 Scope', 2)
bullet([
    'Employment Management: Decentralized platform for job posting, selection, and employment agreements.',
    'Transaction Transparency: Secure, tamper-proof payment records on Solana blockchain.',
    'Target Users: Industrial workers, service workers, and small-to-medium organizations.',
    'Supported Job Categories: Construction, Delivery, Domestic Help, Event Staffing, Agriculture, Cleaning, Security.',
])

h('2.2 Methodology', 2)
bullet([
    'Architecture: Hybrid Web3 — Solana on-chain smart contract + off-chain Node.js API + React frontend.',
    'Blockchain Platform: Solana (Program ID: 4f9fP5Aoz7Tcu7Z5J7WWhTRUa757QnK91JvpM1Zyg7BM).',
    'Smart Contract Language: Rust with Anchor framework.',
    'Backend: Node.js / Express.js with MongoDB (Mongoose ODM).',
    'Frontend: React 18 + Vite + TailwindCSS + Solana Wallet Adapter.',
])

# ─── SECTION 1: CODING / IMPLEMENTATION ──────────────────────────────────────
doc.add_page_break()
h('3. Coding / Implementation')

# i. Architecture & Design
h('i. Architecture & Design', 2)
p('System Architecture Diagram (High-Level):', bold=True)
p('The platform uses a 3-tier hybrid Web3 architecture:')
bullet([
    'Tier 1 – Frontend (React + Vite): Wallet-connected SPA; communicates with both the backend API and Solana RPC.',
    'Tier 2 – Backend (Node.js / Express): Off-chain REST API; stores rich metadata in MongoDB; issues JWT tokens.',
    'Tier 3 – Blockchain (Solana / Anchor): On-chain escrow, job status state machine, and automated fund release.',
])
doc.add_paragraph()

# Architecture table
p('System Architecture Component Table:', bold=True)
styled_table(
    ['Layer', 'Technology', 'Role'],
    [
        ('Frontend', 'React 18, Vite, TailwindCSS, Framer Motion', 'UI for Workers, Employers & Admins'),
        ('Wallet Integration', 'Solana Wallet Adapter (Phantom, Solflare)', 'Web3 Auth & Transaction Signing'),
        ('Backend API', 'Node.js, Express.js, JWT, Zod', 'Off-chain REST API, Auth, Validation'),
        ('Database', 'MongoDB, Mongoose ODM', 'Off-chain profile & job metadata'),
        ('Smart Contract', 'Rust, Anchor Framework', 'On-chain Escrow, Job State Machine'),
        ('Blockchain', 'Solana (Devnet / Localnet)', 'Immutable transaction records'),
        ('Containerization', 'Docker, Docker Compose', 'Service orchestration'),
    ]
)

doc.add_paragraph()
p('Modules / Components Breakdown:', bold=True)
bullet([
    'User Registration Module – Worker & Company profile creation with wallet-based identity.',
    'Job Posting Module – Employer creates job and locks SOL payment into Escrow PDA.',
    'Job Application Module – Workers browse, filter, and apply for open jobs.',
    'Employment Agreement Module – Smart contract governs assignment, proof submission, and payment.',
    'OTP Verification Module – Start/End OTP system confirms on-site attendance.',
    'Dispute Resolution Module – 3-day dispute window; admin arbitration with split/favor resolution.',
    'Admin Panel Module – Admin dashboard for profile verification and dispute management.',
    'Rating & Review Module – Post-job ratings stored on-chain via UserRating account.',
])

# ii. Development Details
h('ii. Development Details', 2)
p('Key Module Explanations:', bold=True)

p('A) Smart Contract (employment_contract.rs) – 6-Step Job Lifecycle:', bold=True)
bullet([
    'Step 1 – post_job(): Employer submits job details; SOL payment immediately transferred to Escrow PDA via CPI.',
    'Step 2 – assign_worker(): Employer selects a worker from applicants; Job moves to InProgress.',
    'Step 3 – submit_proof_of_work(): Worker submits proof (OTP/GPS/Photo); dispute_deadline set to now + 3 days.',
    'Step 4 – release_payment(): After dispute period expires, funds auto-released to worker wallet.',
    'Step 5 – create_dispute(): Employer raises dispute within 3-day window; job moves to Disputed.',
    'Step 6 – resolve_dispute(): Admin resolves with FavorWorker / FavorEmployer / Split outcome.',
])

p('B) Backend API (Node.js / Express):', bold=True)
bullet([
    'authMiddleware: Verifies JWT Bearer tokens; populates req.user with wallet address and user type.',
    'adminAuthMiddleware: Validates JWT and checks wallet against hardcoded admin public key.',
    'jobController: Manages job creation, application approval, OTP generation/verification, proof recording, and fund transfer logging.',
    'adminController: Handles admin login, fetches/verifies worker and company profiles with PDA assignment.',
    'Validation: Zod schemas validate all API inputs; Mongoose schemas enforce DB-level constraints.',
])

p('C) Frontend (React + Vite):', bold=True)
bullet([
    'WalletContext: Global React context storing connected wallet address and verification state.',
    'Role-based routing: Separate dashboards for /worker/dashboard, /company/dashboard, /admin/dashboard.',
    'Solana Wallet Adapter: Phantom, Solflare, Torus wallets; connection to local/devnet RPC.',
    'useAuth hook: Handles wallet signing and backend JWT authentication flow.',
])

p('Database Design (MongoDB Collections):', bold=True)
styled_table(
    ['Collection', 'Key Fields', 'Purpose'],
    [
        ('WorkerProfile', 'walletAddress, name, skills[], jobCategories[], rating, PDAAddress', 'Stores worker off-chain metadata'),
        ('WorkerSkill', 'workerWallet, skillName, category, proficiencyLevel, experienceYears', 'Detailed skill records'),
        ('WorkerAvailability', 'workerWallet, dayOfWeek, startTime, endTime', 'Availability schedule'),
        ('CompanyProfile', 'walletAddress, companyName, companyType, location, rating', 'Employer off-chain metadata'),
        ('Job', 'jobPDA, escrowPDA, companyWallet, title, status, applications[], proofOfWork, disputePeriod', 'Central job document'),
        ('JobApplication', 'jobId, workerWallet, status, coverLetter', 'Worker applications'),
    ]
)

# iii. Code Highlights
doc.add_paragraph()
h('iii. Code Highlights', 2)
p('A) Smart Contract – Escrow Fund Lock on Job Post (Rust):', bold=True)
code_block('// STEP 1: Employer posts job AND locks payment in escrow immediately\npub fn post_job(ctx: Context<PostJob>, payment_amount: u64, ...) -> Result<()> {\n    let cpi_context = CpiContext::new(\n        ctx.accounts.system_program.to_account_info(),\n        anchor_lang::system_program::Transfer {\n            from: ctx.accounts.employer.to_account_info(),\n            to: escrow.to_account_info(),  // PDA escrow\n        },\n    );\n    anchor_lang::system_program::transfer(cpi_context, payment_amount)?;\n    escrow.is_locked = true;\n}')

doc.add_paragraph()
p('B) Backend – JWT Auth Middleware (Node.js):', bold=True)
code_block('// authMiddleware.js\nconst decoded = jwt.verify(token, process.env.JWT_SECRET);\nreq.user = { userId: decoded.userId, walletAddress: decoded.walletAddress,\n             userType: decoded.userType };\nnext();')

doc.add_paragraph()
p('C) Backend – Auto INR Calculation Pre-save Hook (Mongoose):', bold=True)
code_block('// jobModel.js pre-save hook\nif (!this.paymentAmountINR && this.paymentAmount) {\n    const SOL = this.paymentAmount / 1e9;   // lamports to SOL\n    this.paymentAmountINR = Math.round(SOL * 8000);  // ~8000 INR per SOL\n}')

doc.add_paragraph()
p('D) Frontend – Wallet Authentication Flow (React):', bold=True)
code_block("// useAuth.js\nconst signMessage = async () => {\n    const msg = new TextEncoder().encode('Sign in to DeWages Network');\n    const sig = await signMsg(msg);  // Phantom wallet signs\n    await axios.post('/v1/auth/signin', { walletAddress, signature });\n};")

# iv. Tools & Environment
doc.add_paragraph()
h('iv. Tools & Environment', 2)
styled_table(
    ['Tool / Platform', 'Version', 'Purpose'],
    [
        ('VS Code', 'Latest', 'Primary IDE for all development'),
        ('Node.js', 'v18+', 'Backend runtime'),
        ('React + Vite', '18 / 7.x', 'Frontend SPA framework'),
        ('Rust + Anchor', '0.29.0', 'Smart contract development'),
        ('MongoDB', '7.x', 'Off-chain database'),
        ('Docker', 'Latest', 'Containerized deployment'),
        ('Solana CLI', 'Localnet', 'Blockchain test environment'),
        ('Phantom Wallet', 'Browser Extension', 'Web3 wallet for testing'),
        ('Postman', 'Latest', 'API testing and documentation'),
        ('Jest', '29.x', 'Backend automated testing framework'),
    ]
)

# ─── SECTION 2: VALIDATION & TESTING ─────────────────────────────────────────
doc.add_page_break()
h('4. Validation & Testing Section')

# i. Testing Strategy
h('i. Testing Strategy', 2)
p('The backend test suite consists of 147 tests across 6 test suites, achieving 100% pass rate. The testing framework uses Jest with MongoDB Memory Server for isolated, fast, and reproducible tests.')
styled_table(
    ['Type', 'Description', 'Files'],
    [
        ('Unit Testing', 'Tests individual model schemas, Zod validators, and middleware in isolation', 'models.test.js, schemas.test.js, middleware.test.js'),
        ('Integration Testing', 'Tests controller functions with real in-memory MongoDB interactions', 'admin.test.js, job.test.js'),
        ('System Testing', 'End-to-end workflows: full job lifecycle, dispute resolution, multi-user scenarios', 'workflow.test.js'),
        ('User Acceptance Testing', 'Validates business rules: OTP flow, dispute periods, fund transfer blocking, role-based access', 'Covered in system & integration tests'),
    ]
)

# ii. Test Cases
doc.add_paragraph()
h('ii. Test Cases', 2)
p('Sample test cases from the 147-test suite (table format):', bold=True)
styled_table(
    ['Test ID', 'Input', 'Expected Output', 'Actual Output', 'Status'],
    [
        ('TC-M01', 'Valid worker profile data', 'Worker created with all fields saved', 'Worker created, DB constraints passed', 'PASS'),
        ('TC-M02', 'Worker data without walletAddress', 'Mongoose ValidationError thrown', 'ValidationError thrown correctly', 'PASS'),
        ('TC-M28', 'paymentAmount = -100', 'ValidationError (min: 0)', 'ValidationError thrown', 'PASS'),
        ('TC-M30', '0.5 SOL, no INR provided', 'Auto-calc: 0.5 x 8000 = ₹4000', 'paymentAmountINR = 4000', 'PASS'),
        ('TC-M31', '3 applications (mixed status)', 'Virtual: total=3, pending=1, approved=1', 'Virtual fields computed correctly', 'PASS'),
        ('TC-MW01', 'Valid Bearer JWT token', 'next() called, req.user populated', 'Middleware passed, user set', 'PASS'),
        ('TC-MW05', 'Expired JWT token', '401 Token expired', '401 returned correctly', 'PASS'),
        ('TC-MW09', 'Admin wallet token', 'next(), req.user.isAdmin = true', 'Admin access granted', 'PASS'),
        ('TC-MW10', 'Non-admin wallet token', '403 Admin privileges required', '403 returned correctly', 'PASS'),
        ('TC-J09', 'Duplicate worker application', 'Error: Worker already applied', 'Error thrown as expected', 'PASS'),
        ('TC-J21', '3-day dispute period set', 'isActive=true, 3-day diff verified', 'Dispute period stored correctly', 'PASS'),
        ('TC-J22', 'Expired dispute period', 'isExpired=true, status=completed', 'Auto-completion logic verified', 'PASS'),
        ('TC-A01', 'Valid admin credentials', '200 OK + JWT token returned', 'Token returned', 'PASS'),
        ('TC-A09', 'Verify worker (PDA + isVerified)', '200, DB state updated', 'Worker verified in DB', 'PASS'),
        ('TC-SYS01', 'Full 14-step job lifecycle', 'Job completed + funds transferred', 'All lifecycle steps passed', 'PASS'),
        ('TC-SYS02', 'Employer raises dispute in window', 'Status = disputed, dispute stored', 'Dispute handled correctly', 'PASS'),
        ('TC-SYS10', 'Fund transfer during active dispute', 'Transfer blocked', 'canTransferFunds = false', 'PASS'),
    ]
)

# iii. Validation Techniques
doc.add_paragraph()
h('iii. Validation Techniques', 2)
p('How Correctness Was Verified:', bold=True)
styled_table(
    ['Technique', 'Description'],
    [
        ('Schema Validation', 'Mongoose schema constraints: required fields, enum values, min/max ranges, maxlength, unique indexes.'),
        ('Zod Input Validation', 'All API payloads validated via Zod schemas before reaching controllers. Rejects malformed inputs.'),
        ('JWT Token Verification', 'Valid, expired, invalid, and wrong-secret tokens tested. Admin wallet checked against hardcoded pubkey.'),
        ('Database State Assertions', 'After every operation, MongoDB queried to verify persisted state matches expected values.'),
        ('Virtual Field Computation', 'Mongoose virtual fields (isFullyVerified, totalApplications) verified for correctness.'),
        ('Pre-save Hook Verification', 'Auto-calculations (INR conversion, updatedAt timestamps) confirmed on every save.'),
        ('Business Logic Validation', 'Domain rules tested: duplicate application prevention, status transitions, OTP expiry, fund lock enforcement.'),
        ('Boundary Testing', 'Skills max=20, rating 1-5, limit max=100, page min=1, bio maxlength=500 all verified.'),
    ]
)

doc.add_paragraph()
p('Data Validation Methods:', bold=True)
styled_table(
    ['Method', 'Applied To'],
    [
        ('Required Field Validation', 'walletAddress (Worker/Company), jobPDA, escrowPDA, title, description, category'),
        ('Enum Constraints', 'experienceLevel, jobCategories, companyType, status, proofType, disputeStatus'),
        ('Range / Boundary Validation', 'rating (0-5), skills (max 20), jobCategories (max 5), pagination (page>=1, limit<=100)'),
        ('Format Validation', 'Email (regex), phone (regex), Solana public key (Base58 format)'),
        ('Uniqueness Constraints', 'jobPDA (unique index), JobApplication (jobId + workerWallet composite unique)'),
        ('Default Value Verification', 'All model defaults verified: rating=0, isActive=false, status=open, fundTransfer.isTransferred=false'),
    ]
)

# iv. Results & Analysis
doc.add_paragraph()
h('iv. Results & Analysis', 2)
p('Test Execution Summary:', bold=True)
styled_table(
    ['Test Suite', 'Type', 'Tests', 'Passed', 'Failed', 'Pass Rate'],
    [
        ('unit/models.test.js', 'Unit', '39', '39', '0', '100%'),
        ('unit/schemas.test.js', 'Unit', '36', '36', '0', '100%'),
        ('unit/middleware.test.js', 'Unit', '13', '13', '0', '100%'),
        ('integration/admin.test.js', 'Integration', '15', '15', '0', '100%'),
        ('integration/job.test.js', 'Integration', '33', '33', '0', '100%'),
        ('system/workflow.test.js', 'System', '11', '11', '0', '100%'),
        ('TOTAL', '–', '147', '147', '0', '100%'),
    ]
)
doc.add_paragraph()
p('Key Observations:', bold=True)
bullet([
    'All 147 tests pass consistently with zero flaky tests across 6 suites.',
    'MongoDB Memory Server provides isolated, reproducible test runs with no external DB dependency.',
    'Complete job lifecycle tested end-to-end: creation → escrow lock → OTP → proof → dispute → fund release.',
    'Validation tested at two layers: Zod (API input) and Mongoose (DB persistence).',
    'Admin role-based access control verified with hardcoded Solana public key comparison.',
    'Execution time: ~4.5 seconds for all 147 tests — fast enough for CI/CD integration.',
])

# ─── RESULTS ─────────────────────────────────────────────────────────────────
doc.add_page_break()
h('5. Results')
p('The Blockchain Based Employment Platform provides secure and transparent employment management for industrial and service workforce. It enables direct interaction between workers and employers and ensures timely payments using smart contracts. The system securely stores employment records on blockchain and improves trust and efficiency.')
p('The smart contract enforces a strict job state machine (Open → InProgress → PendingVerification → Completed/Disputed) with funds locked in a Program Derived Address (PDA) escrow throughout. The 3-day dispute window provides adequate time for employer review without unnecessarily delaying worker payments. The automated test suite of 147 tests achieving 100% pass rate demonstrates the robustness and reliability of the platform.')

# ─── CONCLUSION ──────────────────────────────────────────────────────────────
h('6. Conclusion')
p('The Blockchain Based Employment Platform provides a secure and transparent solution for managing employment in industrial and service sectors. By using blockchain and smart contracts on the Solana network, the system ensures fair job management, timely payments, and reliable record keeping. The platform improves trust between workers and employers and enhances overall efficiency.')
p('The hybrid architecture—combining the speed and richness of a traditional backend with the trustlessness of a smart contract—strikes the right balance for a production-grade decentralized employment platform. Future enhancements may include mobile wallet support, IPFS-based document storage for worker credentials, and on-chain reputation scoring.')

# ─── REFERENCES ──────────────────────────────────────────────────────────────
h('7. References')
refs = [
    'Solana Developer Documentation – https://docs.solana.com',
    'Anchor Framework Documentation – https://www.anchor-lang.com',
    'Mongoose ODM Documentation – https://mongoosejs.com',
    'Solana Wallet Adapter – https://github.com/solana-labs/wallet-adapter',
    'React 18 Documentation – https://react.dev',
    'Jest Testing Framework – https://jestjs.io',
    'MongoDB Memory Server – https://github.com/nodkz/mongodb-memory-server',
    'Zod Schema Validation – https://zod.dev',
]
for i, r in enumerate(refs):
    doc.add_paragraph(f'[{i+1}] {r}', style='List Number')

doc.save('/Users/tejasshinde/BigProjects/dewages-network/DeWages_Network_Documentation.docx')
print('Document saved successfully.')
